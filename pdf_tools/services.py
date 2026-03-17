import io
import logging
from django.core.files.base import ContentFile

logger = logging.getLogger("unitools")


def _pypdf2():
    try:
        from pypdf import PdfReader, PdfWriter

        return PdfReader, PdfWriter
    except (ModuleNotFoundError, ImportError):
        try:
            from PyPDF2 import PdfReader, PdfWriter

            logger.warning("Using legacy PyPDF2 backend. Install pypdf for best compatibility.")
            return PdfReader, PdfWriter
        except ModuleNotFoundError as inner_exc:
            logger.exception("PDF dependency missing: %s", inner_exc)
            raise RuntimeError("PDF tools require pypdf (or PyPDF2). Install dependencies from requirements.txt.") from inner_exc


def merge_pdfs(file_list):
    try:
        PdfReader, PdfWriter = _pypdf2()
        writer = PdfWriter()
        for file_obj in file_list:
            reader = PdfReader(file_obj)
            for page in reader.pages:
                writer.add_page(page)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        return ContentFile(buf.read(), name="merged.pdf")
    except Exception as exc:
        logger.exception("merge_pdfs failed: %s", exc)
        raise


def compress_pdf(file_obj, target_bytes=None):
    try:
        PdfReader, PdfWriter = _pypdf2()

        reader = PdfReader(file_obj)
        writer = PdfWriter()
        for page in reader.pages:
            page.compress_content_streams()
            writer.add_page(page)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        output = buf.read()
        result_name = "compressed.pdf"
        if target_bytes:
            actual = len(output)
            target_kb = max(1, int(target_bytes / 1024))
            actual_kb = max(1, int(actual / 1024))
            result_name = f"compressed_target-{target_kb}kb_actual-{actual_kb}kb.pdf"
        return ContentFile(output, name=result_name)
    except Exception as exc:
        logger.exception("compress_pdf failed: %s", exc)
        raise


def split_pdf(file_obj, start_page=1, end_page=1):
    try:
        PdfReader, PdfWriter = _pypdf2()

        reader = PdfReader(file_obj)
        writer = PdfWriter()
        total = len(reader.pages)
        start_idx = max(0, start_page - 1)
        end_idx = min(total, end_page)
        for page_index in range(start_idx, end_idx):
            writer.add_page(reader.pages[page_index])
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        return ContentFile(buf.read(), name=f"split_{start_page}_{end_page}.pdf")
    except Exception as exc:
        logger.exception("split_pdf failed: %s", exc)
        raise


def rotate_pdf(file_obj, degrees=90):
    try:
        PdfReader, PdfWriter = _pypdf2()

        reader = PdfReader(file_obj)
        writer = PdfWriter()
        for page in reader.pages:
            page.rotate(degrees)
            writer.add_page(page)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        return ContentFile(buf.read(), name=f"rotated_{degrees}.pdf")
    except Exception as exc:
        logger.exception("rotate_pdf failed: %s", exc)
        raise


def add_watermark(file_obj, watermark_text="CONFIDENTIAL"):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ModuleNotFoundError as exc:
        raise RuntimeError("Watermark requires reportlab.") from exc
    try:
        PdfReader, PdfWriter = _pypdf2()
        wm_buffer = io.BytesIO()
        can = canvas.Canvas(wm_buffer, pagesize=letter)
        can.setFont("Helvetica", 44)
        can.setFillColorRGB(0.7, 0.7, 0.7)
        if hasattr(can, "setFillAlpha"):
            can.setFillAlpha(0.3)
        can.saveState()
        can.translate(280, 380)
        can.rotate(45)
        can.drawCentredString(0, 0, watermark_text)
        can.restoreState()
        can.save()
        wm_buffer.seek(0)

        wm_reader = PdfReader(wm_buffer)
        wm_page = wm_reader.pages[0]
        source = PdfReader(file_obj)
        writer = PdfWriter()
        for page in source.pages:
            page.merge_page(wm_page)
            writer.add_page(page)
        out = io.BytesIO()
        writer.write(out)
        out.seek(0)
        return ContentFile(out.read(), name="watermarked.pdf")
    except Exception as exc:
        logger.exception("add_watermark failed: %s", exc)
        raise


def pdf_to_word(file_obj):
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise RuntimeError("PDF to Word requires python-docx.") from exc
    try:
        PdfReader, _ = _pypdf2()
        reader = PdfReader(file_obj)
        document = Document()
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                document.add_paragraph(text)
        out = io.BytesIO()
        document.save(out)
        out.seek(0)
        return ContentFile(out.read(), name="converted.docx")
    except Exception as exc:
        logger.exception("pdf_to_word failed: %s", exc)
        raise


def word_to_pdf(file_obj):
    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ModuleNotFoundError as exc:
        raise RuntimeError("Word to PDF requires python-docx and reportlab.") from exc
    try:
        doc = Document(file_obj)
        out = io.BytesIO()
        pdf = SimpleDocTemplate(out, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                story.append(Paragraph(paragraph.text, styles["Normal"]))
                story.append(Spacer(1, 8))
        pdf.build(story)
        out.seek(0)
        return ContentFile(out.read(), name="converted.pdf")
    except Exception as exc:
        logger.exception("word_to_pdf failed: %s", exc)
        raise
