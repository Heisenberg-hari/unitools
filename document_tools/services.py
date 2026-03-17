import logging
import difflib
import io
import re
from collections import Counter
from core.llm import call_llm, is_llm_enabled

logger = logging.getLogger("unitools")


def summarize_text(text, max_length=1200, use_llm=True):
    try:
        clean = " ".join(text.split())
        if not clean:
            return ""
        if use_llm and is_llm_enabled():
            prompt = (
                "Summarize the following document text clearly with bullet points for key points, "
                "important numbers, and action items. Keep it complete but concise.\n\n"
                f"{clean[:12000]}"
            )
            try:
                return call_llm("You are an expert document summarizer.", prompt, max_output_tokens=1200)
            except Exception as llm_exc:
                logger.warning("LLM summarize fallback: %s", llm_exc)
        return clean[:max_length].strip() + ("..." if len(clean) > max_length else "")
    except Exception as exc:
        logger.exception("summarize_text failed: %s", exc)
        raise


def compare_text(text_a, text_b, use_llm=True):
    try:
        a_lines = text_a.splitlines()
        b_lines = text_b.splitlines()
        diff = "\n".join(difflib.unified_diff(a_lines, b_lines, fromfile="A", tofile="B", lineterm=""))
        explanation = ""
        if use_llm and is_llm_enabled():
            prompt = (
                "Compare Document A and Document B and provide a concise summary of differences: "
                "added topics, removed topics, changed facts, and risk-impact.\n\n"
                f"Document A:\n{text_a[:8000]}\n\nDocument B:\n{text_b[:8000]}"
            )
            try:
                explanation = call_llm("You are a precise document comparison assistant.", prompt, max_output_tokens=900)
            except Exception as llm_exc:
                logger.warning("LLM compare fallback: %s", llm_exc)
        return {"diff": diff, "explanation": explanation}
    except Exception as exc:
        logger.exception("compare_text failed: %s", exc)
        raise


def extract_text_from_upload(file_obj):
    try:
        name = (getattr(file_obj, "name", "") or "").lower()
        ext = name.rsplit(".", 1)[-1] if "." in name else ""
        content = file_obj.read()
        file_obj.seek(0)

        if ext in ("txt", "md", "csv", "log"):
            return content.decode("utf-8", errors="replace")

        if ext == "pdf":
            try:
                from pypdf import PdfReader
            except ModuleNotFoundError:
                try:
                    from PyPDF2 import PdfReader
                except ModuleNotFoundError as exc:
                    raise RuntimeError("PDF parsing requires pypdf (or PyPDF2).") from exc
            reader = PdfReader(io.BytesIO(content))
            return "\n".join((page.extract_text() or "") for page in reader.pages)

        if ext in ("docx",):
            try:
                from docx import Document
            except ModuleNotFoundError as exc:
                raise RuntimeError("DOCX parsing requires python-docx.") from exc
            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        if ext == "doc":
            # Legacy DOC is binary; best effort decode fallback.
            return content.decode("utf-8", errors="replace")

        return content.decode("utf-8", errors="replace")
    except Exception as exc:
        logger.exception("extract_text_from_upload failed: %s", exc)
        raise


def translate_text(text, target_lang="es"):
    clean = (text or "").strip()
    if not clean:
        return ""
    if is_llm_enabled():
        prompt = (
            f"Translate the following text to language code '{target_lang}'. "
            "Preserve names, numbers, and formatting.\n\n"
            f"{clean[:12000]}"
        )
        try:
            return call_llm("You are a professional translator.", prompt, max_output_tokens=1200)
        except Exception as llm_exc:
            logger.warning("LLM translate fallback: %s", llm_exc)
    try:
        from deep_translator import GoogleTranslator
    except ModuleNotFoundError:
        logger.warning("translate_text fallback to source text: deep-translator is not installed.")
        return clean
    try:
        result = GoogleTranslator(source="auto", target=target_lang).translate(clean)
        return result or clean
    except Exception as exc:
        logger.warning("translate_text fallback to source text: %s", exc)
        return clean


def analyze_text(text):
    clean = (text or "").strip()
    if not clean:
        return {
            "characters": 0,
            "characters_no_spaces": 0,
            "words": 0,
            "sentences": 0,
            "paragraphs": 0,
            "avg_word_length": 0,
            "reading_time_minutes": 0,
            "top_words": [],
        }
    words = re.findall(r"[A-Za-z0-9']+", clean.lower())
    sentences = [s for s in re.split(r"[.!?]+", clean) if s.strip()]
    paragraphs = [p for p in re.split(r"\n\s*\n", clean) if p.strip()]
    top_words = Counter(w for w in words if len(w) > 2).most_common(10)
    avg_word_length = round(sum(len(w) for w in words) / len(words), 2) if words else 0
    reading_time = round(len(words) / 200, 2) if words else 0
    return {
        "characters": len(clean),
        "characters_no_spaces": len(clean.replace(" ", "")),
        "words": len(words),
        "sentences": len(sentences),
        "paragraphs": len(paragraphs),
        "avg_word_length": avg_word_length,
        "reading_time_minutes": reading_time,
        "top_words": top_words,
    }


def docx_to_pdf(file_obj):
    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ModuleNotFoundError as exc:
        raise RuntimeError("DOCX to PDF requires python-docx and reportlab.") from exc
    try:
        doc = Document(file_obj)
        buf = io.BytesIO()
        pdf = SimpleDocTemplate(buf, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        for para in doc.paragraphs:
            if para.text.strip():
                story.append(Paragraph(para.text, styles["Normal"]))
                story.append(Spacer(1, 8))
        pdf.build(story)
        buf.seek(0)
        return buf
    except Exception as exc:
        logger.exception("docx_to_pdf failed: %s", exc)
        raise
